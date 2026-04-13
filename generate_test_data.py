# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "faker",
#     "python-docx",
#     "fpdf2",
#     "requests",
# ]
# ///

import json
import os
import random
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import re
from faker import Faker
import uuid

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "qwen3:4b"

fake = Faker()


def generate_text_with_ollama(prompt):
    """Hits the local Ollama instance to generate creative unstructured text."""
    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": MODEL_NAME,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.95},
            },
        )
        response.raise_for_status()
        return response.json().get("response", "").strip()
    except Exception as e:
        print(f"[!] Ollama connection failed. Is it running? {e}")
        return f"Dummy generated content because Ollama failed: {fake.paragraph()}"


def create_job_description(job_id, role_title):
    """Generates a rich, non-technical and technical mixed job description."""
    print(f"Generating rich job description for {role_title}...")
    company_name = fake.company()
    industry = fake.bs()

    prompt = f"""
Write a realistic, concise, and detailed job description for a "{role_title}" at a fictional company named "{company_name}" in the {industry} industry.

Rules:
- Output plain text only. No markdown code fences.
- Keep it realistic, specific, and internally consistent.
- Match the responsibilities and requirements to the role seniority and the industry.
- Avoid generic filler, hype, and repetition.
- Include concrete tools, technologies, and expectations where relevant.
- Make it sound like a real hiring post from a modern company.

Use exactly this structure and headings:

About {company_name}
Write 2-3 sentences about the company, product/service, mission, and industry context.

Working Conditions & Perks
List 4-6 items including work model (onsite/hybrid/remote), flexibility, compensation/perks, learning budget, and benefits.

Requirements (Must-Have)
List 4-6 specific requirements focused on core skills, years of experience, and essential tools/technologies.

Nice-to-Have
List 2-4 bonus skills or experiences that are helpful but not required.

Responsibilities
List 4-6 concrete day-to-day responsibilities for this role.

Important:
- Requirements must be testable and role-relevant.
- Responsibilities must reflect the actual work this role would do.
- Do not mention that the company is fabricated, fictional, or imaginary.
"""

    description_text = generate_text_with_ollama(prompt)

    return {"id": job_id, "title": role_title, "description": description_text}


def create_candidate_cv_text(role_title, match_type):
    """Generates a CV matching (or completely failing) the role with varied formats."""
    print(f"Generating a {match_type} candidate CV for {role_title}...")

    candidate_name = fake.name()
    candidate_email = fake.email()
    candidate_phone = fake.phone_number()

    if match_type == "strong":
        prompt_instruction = (
            f"Create a strong candidate for {role_title}: "
            f"highly relevant background, 5+ years of experience, solid achievements, "
            f"and most of the core tools/skills expected for the role."
        )
    else:
        prompt_instruction = (
            f"Create a weak-but-plausible candidate for {role_title}: "
            f"same or adjacent field, but clearly less qualified than a strong candidate. "
            f"Use 0-2 years of experience, smaller scope, weaker achievements, and only some of the core tools/skills. "
            f"Do NOT make them from a completely unrelated profession."
        )

    layouts = [
        "CLASSIC: Contact at the top, then a short Summary, Experience (bullet points), Education, and list Skills at the end.",
        "MODERN: Contact at the top, followed immediately by a huge block of Technical Skills, then Work Experience, and minimal Education.",
        "ACADEMIC: Put Education and Research/Projects first, followed by Work Experience.",
        "NARRATIVE: Do not use bullet points anywhere. Write the entire work experience and skills as flowing narrative paragraphs instead of lists.",
    ]
    chosen_layout = random.choice(layouts)

    prompt = f"""
    Write a realistic resume text for a candidate.
    {prompt_instruction}
    
    Candidate Details REQUIRED:
    - Name: {candidate_name}
    - Email: {candidate_email}
    - Phone: {candidate_phone}
    
    You MUST apply this layout/format constraint:
    "{chosen_layout}"
    
    Include: The exact Name, Email, and Phone provided above, followed by a Summary, Work Experience, Education, and Skills according to the layout constraint.
    Use **double asterisks** for bolding headers or keywords (like **Experience** or **Python**). Do not use markdown code blocks.
    """
    return generate_text_with_ollama(prompt)


def save_as_pdf(text, filename):
    pdf = FPDF()
    pdf.add_page()
    margin = random.choice([10, 11, 13, 15])
    pdf.set_auto_page_break(auto=True, margin=margin)

    font_choice = random.choice(["Helvetica", "Times", "Courier"])
    size_choice = random.choice([10, 11, 12])
    pdf.set_font(font_choice, size=size_choice)

    safe_text = text.encode("latin-1", "replace").decode("latin-1")

    try:

        pdf.multi_cell(0, 6, txt=safe_text, markdown=True)
    except Exception:

        safe_text = safe_text.replace("**", "")
        pdf.multi_cell(0, 6, txt=safe_text)

    pdf.output(filename)


def save_as_docx(text, filename):
    doc = Document()

    font_name = random.choice(["Arial", "Times New Roman", "Calibri", "Cambria"])
    doc.styles["Normal"].font.name = font_name

    align_headers_center = random.choice([True, False])

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()

        if align_headers_center and len(line) < 40 and line.isupper():
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    doc.save(filename)


def main():

    roles = [
        {"id": "python-backend-dev", "title": "Backend Developer (Python)"},
        {"id": "frontend-react-dev", "title": "Frontend Developer (React)"},
        {"id": "data-analyst", "title": "Data Analyst"},
        {"id": "data-scientist", "title": "Data Scientist"},
        {"id": "ai-engineer", "title": "AI Engineer"},
        {"id": "sr-data-scientist", "title": "Senior Data Scientist"},
        {"id": "jr-data-scientist", "title": "Jr. Data Scientist"},
        {"id": "ml-engineer", "title": "Machine Learning Engineer"},
        {"id": "devops-engineer", "title": "DevOps Engineer"},
        {"id": "product-manager", "title": "Product Manager"},
        {"id": "ux-designer", "title": "UX Designer"},
        {"id": "qa-engineer", "title": "Jr. QA Engineer"},
        {"id": "fs-developer", "title": "Full Stack Developer"},
    ]

    jobs_json_data = []
    output_dir = "test_cvs"
    os.makedirs(output_dir, exist_ok=True)

    print("Starting HR Ops Mock Data Generation")

    for role in roles:
        job_data = create_job_description(role["id"], role["title"])
        jobs_json_data.append(job_data)

        strong_cv_text = create_candidate_cv_text(role["title"], "strong")
        weak_cv_text = create_candidate_cv_text(role["title"], "weak")

        strong_ext = random.choice(["pdf", "docx"])
        weak_ext = random.choice(["pdf", "docx"])

        strong_filename = os.path.join(
            output_dir, f"{role['id']}_{uuid.uuid4().hex[:8]}_StrongMatch.{strong_ext}"
        )
        weak_filename = os.path.join(
            output_dir, f"{role['id']}_{uuid.uuid4().hex[:8]}_WeakMatch.{weak_ext}"
        )

        if strong_ext == "pdf":
            save_as_pdf(strong_cv_text, strong_filename)
        else:
            save_as_docx(strong_cv_text, strong_filename)

        if weak_ext == "pdf":
            save_as_pdf(weak_cv_text, weak_filename)
        else:
            save_as_docx(weak_cv_text, weak_filename)
        print(f"-> Saved: {strong_filename}")
        print(f"-> Saved: {weak_filename}")

    with open("jobs.json", "w", encoding="utf-8") as f:
        json.dump(jobs_json_data, f, indent=2)
    print("\n ✅ Successfully updated jobs.json and generated CVs in test_cvs")


if __name__ == "__main__":
    main()
